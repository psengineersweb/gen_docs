import streamlit as st
import requests
import pandas as pd
from io import BytesIO
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
import phpserialize

import json
post_id=0

# Key Mapping for readability
def add_images_to_document(image_urls, doc):
    """
    Adds images from the given URLs to an existing Word document object.

    :param image_urls: List of image URLs.
    :param doc: An existing `Document` object.
    :return: Updated `Document` object.
    """
    for url in image_urls:
        try:
            response = requests.get(url)
            response.raise_for_status()  # Ensure the request was successful
            image_stream = BytesIO(response.content)

            doc.add_picture(image_stream, width=Inches(4))
            doc.add_paragraph(url)  # Add image URL for reference

        except requests.exceptions.RequestException as e:
            doc.add_paragraph(f"Failed to load image: {url} \nError: {e}")

    return doc  # Return the updated document object


def apply_header_footer(doc):
    """Applies the same header and footer to all sections of the document."""
    for section in doc.sections:
        section.start_type = 2  # Ensure new sections start on a new page

        # === HEADER ===
        header = section.header
        header.is_linked_to_previous = False
        
        # Create a table with a specified width inside the header
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        header_table.allow_autofit = False

        # Set explicit column widths
        cell_text = header_table.cell(0, 0)
        cell_logo = header_table.cell(0, 1)
        cell_text.width = Inches(5)  # 75% width
        cell_logo.width = Inches(1.5)  # 25% width

        # HEADER TEXT
        title = cell_text.paragraphs[0]
        title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = title.add_run("KOUSHIK KUMAR DAS")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue text

        details = cell_text.add_paragraph()
        details.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        text = (
            "M.SC. REAL ESTATE VALUATION, M.I.S.(VALUATION-SURVEYING), AMISE (CIVIL ENGG), "
            "LIFE ASSOCIATE OF THE INSTITUTION OF SURVEYORS (VALUATION - SURVEYING), F.I.V., "
            "APPROVED VALUER, INSTITUTION OF VALUERS, BANKS, INSURANCE, FINANCIAL AND INDUSTRIAL "
            "CORPORATIONS & APPROVED VALUER (CAT-1) OF IMMOVABLE PROPERTY & ENGINEER COMMISSIONER’ "
            "OF THE OFFICE OF THE CITY CIVIL COURT, CALCUTTA."
        )
        run = details.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 255)

        # HEADER LOGO
        paragraph_logo = cell_logo.paragraphs[0]
        paragraph_logo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run_logo = paragraph_logo.add_run()

        try:
            run_logo.add_picture("logo.png", width=Inches(1.5))
        except FileNotFoundError:
            print("Warning: Logo file not found. Add 'logo.png' to the working directory.")

        # === FOOTER ===
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align the footer

        footer_text = (
            "OFFICE: 20/1, CHETLA HAT ROAD, P.O. ALIPORE (H.O.), KOLKATA – 700027\n"
            "MOBILE NOS: 9903068614, 9830776121, 9433240397,\n"
            "E-MAIL- kkdareport@gmail.com valuerkoushik@yahoo.co.in"
        )

        run_footer = footer_paragraph.add_run(footer_text)
        run_footer.italic = True
        run_footer.font.size = Pt(10)
        run_footer.font.color.rgb = RGBColor(0, 0, 255)  # Blue color


key_map = {
    # Part A - Valuation of Land
    "size_of_plot": "Size of Plot",
    "north": "North",
    "south": "South",
    "east": "East",
    "west": "West",
    "total_extent_of_the_plot": "Total Extent of the Plot",
    "prevailing_market_rate": "Prevailing Market Rate",
    "guideline_rate_obtained_from_the_registrars_office": "Guideline Rate",
    "assessedadopted_rate_of_valuation": "Assessed/Adopted Rate",
    "estimated_value_of_land": "Estimated Value of Land",

    # Part C - Extra Items
    "portico": "Portico",
    "ornamental_front_door": "Ornamental Front Door",
    "sit_out_verandah_with_steel_grills": "Sit Out/Verandah with Steel Grills",
    "overhead_water_tank": "Overhead Water Tank",
    "extra_steel_collapsible_gates": "Extra Steel/Collapsible Gates",

    # Part D - Amenities
    "wardrobes": "Wardrobes",
    "glazed_tiles": "Glazed Tiles",
    "extra_sinks_and_bath_tub": "Extra Sinks and Bath Tub",
    "marble__ceramic_tiles_flooring": "Interior Decorations",
    "interior_decorations": "Interior Decorations",
    "architectural_elevation_works": "Architectural Elevation Works",
    "panelling_works": "Panelling Works",
    "aluminium_works": "Aluminium Works",
    "aluminium_hand_rails": "Aluminium Hand Rails",
    "false_ceiling": "False Ceiling",

    # Part E - Miscellaneous
    "separate_toilet_room": "Separate Toilet Room",
    "separate_lumber_room": "Separate Lumber Room",
    "separate_water_tank_sump": "Separate Water Tank/Sump",
    "trees_gardening": "Trees, Gardening",

    # Part F - Services
    "water_supply_arrangements": "Water Supply Arrangements",
    "drainage_arrangements": "Drainage Arrangements",
    "compound_wall": "Compound Wall",
    "c_b_deposits_fittings_etc": "C. B. Deposits, Fittings etc.",
    "pavement": "Pavement",

     "purpose_for_which_the_valuation_is_made": "Purpose for which the valuation is made",
    "date_of_inspection": "Date of inspection",
    "date_on_which_the_valuation_is_made": "Date on which the valuation is made",
    "list_of_documents_produced_for_perusal": "List of documents produced for perusal",
    "name_and_contact_details_of_the_owner": "Name and Contact Details of the owner",
    "brief_description_of_the_property": "Brief description of the property",
    "plot_no__survey_no": "Plot No. / Survey No.",
    "door_no_": "Door No.",
    "t_s_no__village_mouza": "T. S. No. / Village/ Mouza",
    "pin_no": "Pin No:",
    "postal_address_of_the_property": "Postal address of the property",
    "city__town_": "Yes",
    "residential_area_": "Residential Area",
    "commercial_area_": "Commercial Area",
    "classification_of_the_area_": "Classification of the area",
    "coming_under_corporation_limit__village_panchayet__municipality": "Coming under Corporation limit / Village Panchayet / Municipality",
    "covered_under_govt_enactments_ulc_agencyscheduledcantonment_area": "Covered under Govt. Enactments (ULC, Agency/Scheduled/Cantonment Area)",
    "agricultural_land_conversion_to_house_site_planned": "Agricultural Land: Conversion to House Site Planned?",
    "boundaries_of_the_property": "Boundaries of the property",
    "dimensions_of_the_site": "Dimensions of the site",
    "latitude": "Latitude",
    "valuation_least_of_14_a_amp_14_b_": "Longitude",
    "extent_of_the_site_considered_for_valuation": "Extent of the site considered for valuation (least of 14 A & 14 B)",
    "extent_of_the_site": "Extent of the site",
    "whether_occupied_by_": "Whether occupied by the owner / tenant? If occupied by tenant, since how long? Rent received per month.",
    "technical_details_of_the_building": "Technical details of the building",
    "exterior_condition_of_the_": "Exterior Condition of the building",
    "interior_condition_of_the_": "Interior Condition of the building",
    "date_of_issue_and_validity_of_layout_of_approved_": "Date of issue and validity of layout of approved map / plan",
    "approved_map__plan_issuing_": "Approved map / plan issuing authority",
    "whether_genuineness_or_authenticity_of_approved_map__plan_is": "Whether genuineness or authenticity of approved map / plan is verified",
    "any_other_comments_by_our_empanelled_valuers_on_": "Any other comments by our empanelled valuers on authenticity of approved plan"

}

# Function to apply table style with borders
from docx.oxml import OxmlElement, ns

def apply_table_style(table):
    """Apply borders to a table in a Word document."""
    tbl = table._element

    # Correctly find `w:tblPr` using the namespace
    tblPr = tbl.find(ns.qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Remove existing `w:tblBorders` if present
    tblBorders = tblPr.find(ns.qn("w:tblBorders"))
    if tblBorders is not None:
        tblPr.remove(tblBorders)

    # Create new `w:tblBorders`
    tblBorders = OxmlElement("w:tblBorders")

    # Define borders for each side
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(ns.qn("w:val"), "single")  # Solid border
        border.set(ns.qn("w:sz"), "8")  # Thickness
        border.set(ns.qn("w:space"), "0")  # No spacing
        border.set(ns.qn("w:color"), "000000")  # Black color
        tblBorders.append(border)

    # Add the borders to the table properties
    tblPr.append(tblBorders)

image_links = []

def php_to_structured_string(serialized_php_data):
    """
    Converts a PHP-serialized array into a structured string format.
    Extracts image URLs and appends them to the `image_links` list.
    """
    try:
        if isinstance(serialized_php_data, str):
            serialized_php_data = serialized_php_data.encode()  # Convert string to bytes

        # Deserialize PHP array
        parsed_data = phpserialize.loads(serialized_php_data, decode_strings=True)

        # Convert to structured string format
        structured_output = ""
        image_extensions = ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg')

        for index, (key, value) in enumerate(parsed_data.items(), start=1):
            formatted_lines = []
            for k, v in value.items():
                formatted_lines.append(f"  {k.replace('_', ' ').title()}: {v}")
                
                # Check if value is an image URL and append to the global list
                if isinstance(v, str) and v.lower().endswith(image_extensions):
                    image_links.append(v)

            structured_output += "\n".join(formatted_lines) + "\n\n"  # Add spacing between items

        return structured_output.strip()

    except Exception as e:
        return f"Error parsing PHP-serialized data: {e}"

    except Exception:
        return serialized_php_data.decode() if isinstance(serialized_php_data, bytes) else serialized_php_data  
        # Return original value as a fallback

def api_to_dataframe(api_url):
    excluded_keys = {
        "ID", "post_author", "post_date", "post_date_gmt", "post_content", "post_title",
        "post_excerpt", "post_status", "comment_status", "ping_status", "post_password",
        "post_name", "to_ping", "pinged", "post_modified", "post_modified_gmt",
        "post_content_filtered", "post_parent", "guid", "menu_order", "post_type",
        "post_mime_type", "comment_count", "filter", "meta_ID", "object_ID", "_ID", 
        "created", "rel_id", "parent_rel", "parent_object_id", "child_object_id",
    }   
    try:
        response = requests.get(api_url)
        response.raise_for_status()
        data = response.json()
        
        if not data:
            st.warning(f"No data found for {api_url}")
            return [pd.DataFrame(columns=['Key', 'Value']),0]

        if isinstance(data, dict):
            filtered_data = {key_map.get(k, k): v for k, v in data.items() if k not in excluded_keys}
            df = pd.DataFrame(filtered_data.items(), columns=['Key', 'Value'])
        elif isinstance(data, list):
            flattened_data = []
            for item in data:
                filtered_item = {key_map.get(k, k): v for k, v in item.items() if k not in excluded_keys}
                flattened_data.extend(filtered_item.items())
            df = pd.DataFrame(flattened_data, columns=['Key', 'Value'])
        else:
            raise ValueError("Unsupported JSON structure")

        return [df, flattened_data]
    except requests.RequestException as e:
        st.error(f"API request failed: {e}")
    except ValueError as ve:
        st.error(f"Data processing error: {ve}")
    
    return [pd.DataFrame(columns=['Key', 'Value']),flattened_data]

def api_to_dataframe(api_url):
    """
    Fetches data from an API, filters out excluded keys, and converts PHP-serialized values to structured strings.
    """
    excluded_keys = {
        "ID", "post_author", "post_date", "post_date_gmt", "post_content", "post_title",
        "post_excerpt", "post_status", "comment_status", "ping_status", "post_password",
        "post_name", "to_ping", "pinged", "post_modified", "post_modified_gmt",
        "post_content_filtered", "post_parent", "guid", "menu_order", "post_type",
        "post_mime_type", "comment_count", "filter", "meta_ID", "object_ID", "_ID", 
        "created", "rel_id", "parent_rel", "parent_object_id", "child_object_id",
    }   

    try:
        response = requests.get(api_url)
        response.raise_for_status()
        data = response.json()
        
        if not data:
            st.warning(f"No data found for {api_url}")
            return [pd.DataFrame(columns=['Key', 'Value']), 0]

        flattened_data = []
        if isinstance(data, dict):
            filtered_data = {
                key_map.get(k, k): php_to_structured_string(v) if isinstance(v, str) and v.startswith("a:") else v
                for k, v in data.items() if k not in excluded_keys
            }
            df = pd.DataFrame(filtered_data.items(), columns=['Key', 'Value'])
            flattened_data = list(filtered_data.items())

        elif isinstance(data, list):
            for item in data:
                filtered_item = {
                    key_map.get(k, k): php_to_structured_string(v) if isinstance(v, str) and v.startswith("a:") else v
                    for k, v in item.items() if k not in excluded_keys
                }
                flattened_data.extend(filtered_item.items())
            df = pd.DataFrame(flattened_data, columns=['Key', 'Value'])

        else:
            raise ValueError("Unsupported JSON structure")

        return [df, flattened_data]

    except requests.RequestException as e:
        st.error(f"API request failed: {e}")
    except ValueError as ve:
        st.error(f"Data processing error: {ve}")
    
    return [pd.DataFrame(columns=['Key', 'Value']), flattened_data]



def split_and_format_specifications(specifications_df):
    if specifications_df.empty:
        return pd.DataFrame(columns=['Description', 'Ground Floor', 'Others Floor'])

    def split_key(key):
        parts = key.rsplit('_', 2)
        if len(parts) >= 2:
            desc, floor = parts[0], parts[1]
            return desc.replace('_', ' ').capitalize(), floor.replace('_', ' ').capitalize()
        return key.replace('_', ' ').capitalize(), ""

    specifications_df[['Description', 'Floor']] = specifications_df['Key'].apply(lambda x: pd.Series(split_key(x)))
    
    pivot_df = specifications_df.drop(columns=['Key']).pivot_table(
        index='Description', 
        columns='Floor', 
        values='Value', 
        aggfunc='first'
    ).reset_index()
    
    pivot_df.columns.name = None
    pivot_df = pivot_df.rename(columns={'Ground': 'Ground Floor', 'Others': 'Others Floor'})

    return pivot_df

st.title("Valuation Report Generator")
post_id = st.text_input("Enter Post ID:")

if post_id:
    
    api_urls = {
        
        "Generel": f"https://valuerkkda.in/wp-json/generel/generel/?_post_id={post_id}",
        "Part A - Valuation of Land": f"https://valuerkkda.in/wp-json/part-a/part-a/?_post_id={post_id}",
        "Part - B (Valuation of Building)": f"https://valuerkkda.in/wp-json/part-b/part-b/?_post_id={post_id}",
        "Part C - Extra Items": f"https://valuerkkda.in/wp-json/part-c/part-c/?_post_id={post_id}",
        "Part D - Amenities": f"https://valuerkkda.in/wp-json/part-d/part-d/?_post_id={post_id}",
        "Part E - Miscellaneous": f"https://valuerkkda.in/wp-json/part-e/part-e/?_post_id={post_id}",
        "Part F - Services": f"https://valuerkkda.in/wp-json/get_releted_Part_F/get-releted-part-f-/?_post_id={post_id}",
        "Specifications of Construction": f"https://valuerkkda.in/wp-json/specifications/part-specifications/?_post_id={post_id}",
        "TOTAL ABSTRACT OF THE ENTIRE PROPERTY": f"https://valuerkkda.in/wp-json/abstract/abstract/?_post_id={post_id}",
        "Details of the Valuation": f"https://valuerkkda.in/wp-json/details/details/?_post_id={post_id}"
    }
    # mjson={}
    # for post,url in api_urls.items():
    #     mjson[post]=api_to_dataframe(url)[-1]
    # with open("mjson.json", "w", encoding="utf-8") as file:
    #     json.dump(mjson, file, indent=4, ensure_ascii=False)
    # print("JSON data saved to mjson.json")
    
    document = Document()

    apply_header_footer(document)

    sections = {
        "VALUATION REPORT (IN RESPECT OF LAND)": "VALUATION REPORT (IN RESPECT OF LAND)",
        "Owners": "Owners",
        "General": api_to_dataframe(api_urls["Generel"])[0],
        "Part A - Valuation of Land": api_to_dataframe(api_urls["Part A - Valuation of Land"])[0],
        "Part - B (Valuation of Building)": api_to_dataframe(api_urls["Part - B (Valuation of Building)"])[0],
        "Specifications of Construction": split_and_format_specifications(api_to_dataframe(api_urls["Specifications of Construction"])[0]),
        "Details of Valuation": api_to_dataframe(api_urls["Details of the Valuation"])[0],
        "Part C - Extra Items": api_to_dataframe(api_urls["Part C - Extra Items"])[0],
        "Part D - Amenities": api_to_dataframe(api_urls["Part D - Amenities"])[0],
        "Part E - Miscellaneous": api_to_dataframe(api_urls["Part E - Miscellaneous"])[0],
        "Part F - Services": api_to_dataframe(api_urls["Part F - Services"])[0],
        "TOTAL ABSTRACT OF THE ENTIRE PROPERTY":api_to_dataframe(api_urls["TOTAL ABSTRACT OF THE ENTIRE PROPERTY"])[0],
        "PRESENT VALUE OF SAID PROPERTY": "PRESENT VALUE OF SAID PROPERTY",
        "CERTIFICATE OF STABILITY": "CERTIFICATE OF STABILITY",
        "VETTED ESTIMATE": "VETTED ESTIMATE",
        "Format of undertaking to be submitted by Individuals/ proprietor/ partners/ directors DECLARATION- CUM- UNDERTAKING": "Format of undertaking to be submitted by Individuals/ proprietor/ partners/ directors DECLARATION- CUM- UNDERTAKING",
        "Further, I hereby provide the following information.": "Further, I hereby provide the following information.",
        "MODEL CODE OF CONDUCT FOR VALUERS": """
I,  Koushik Kumar Das,   son	of  Late SudhirRanjan Das,    do hereby solemnly affirm and state that: \n
a.	I am a citizen of India \n
b.	I will not undertake valuation of any assets in which I have a direct or indirect interest or become so interested at any time during a period of three years prior to my appointment as valuer or three years after the valuation of assets was conducted by me  \n
c.	The information furnished in my valuation report dated 23.08.2024 is true and correct to the best of my knowledge and belief and I have made an impartial and true valuation of the property.  \n
d.	My representative personally inspected the property on 20.08.2024. The work is not subcontracted to any other valuer and carried out by myself.  \n
e.	Valuation report is submitted in the format as prescribed by the Bank.  \n
f.	I have not been depanelled/ delisted by any other bank and in case any such depanelment by other banks during my empanelment with you, I will inform you within 3 days of such depanelment.  \n
g.	I have not been removed/dismissed from service/employment earlier  \n
h.	I have not been convicted of any offence and sentenced to a term of imprisonment . \n
i.	I have not been found guilty of misconduct in professional capacity  \n
j.	I have not been declared to be unsound mind  \n
k.	I am not an undercharged bankrupt, or has not applied to be adjudicated as a bankrupt;  \n
l.	I am not an undercharged insolvent  \n
m.	I have not been levied a penalty under section 271J of Income-tax Act, 1961 (43 of 1961) and time limit for filing appeal before Commissioner of Incometax (Appeals) or Income-tax Appellate Tribunal, as the case may be has expired, or such penalty has been confirmed by Income-tax Appellate Tribunal, and five years have not elapsed after levy of such penalty  \n
n.	I have not been convicted of an offence connected with any proceeding under the Income Tax Act 1961, Wealth Tax Act 1957 or Gift Tax Act 1958 and  \n
o.	My PAN Card number/Service Tax number as applicable is: AHAPD5062G  \n
p.	I undertake to keep you informed of any events or happenings which would make me ineligible for empanelment as a valuer \n
q.	I have not concealed or suppressed any material information, facts and records and I have made a complete and full disclosure  \n
r.	I have read the Handbook on Policy, Standards and procedure for Real Estate Valuation, 2011 of the IBA and this report is in conformity to the “Standards” enshrined for valuation in the Part-B of the above handbook to the best of my ability. \n
s.	I have read the International Valuation Standards (IVS) and the report submitted to the Bank for the respective asset class is in conformity to the “Standards” as enshrined for valuation in the IVS in “General Standards” and “Asset Standards” as applicable  \n
t.	I abide by the Model Code of Conduct for empanelment of valuer in the Bank. (Annexure V- A signed copy of same to be taken and kept along with this declaration)  \n
u.	I am registered under Section 34 AB of the Wealth Tax Act, 1957. (Strike off, if not applicable)  \n
v.	I am valuer registered with Insolvency & Bankruptcy Board of India (IBBI) (Strike off, if not applicable)  \n
w.	My CIBIL Score and credit worthiness is as per Bank’s guidelines.  \n
x.	I am the proprietor / partner / authorized official of the firm / company, who is competent to sign this valuation report.  \n
y.	I will undertake the valuation work on receipt of Letter of Engagement generated from the system (i.e. LLMS/LOS) only.  \n
z.	Further, I hereby provide the following information.  \n

        """,
    }

    

    for section, content in sections.items():
        st.subheader(section)
        document.add_heading(section, level=2)

        # Check if content is a DataFrame
        if isinstance(content, pd.DataFrame):
            if not content.empty:
                st.table(content)

                # Create table in docx
                docx_table = document.add_table(rows=1, cols=len(content.columns))

                # Apply table style if function exists
                if "apply_table_style" in globals():
                    apply_table_style(docx_table)

                hdr_cells = docx_table.rows[0].cells
                for i, col_name in enumerate(content.columns):
                    hdr_cells[i].text = col_name
                    hdr_cells[i].paragraphs[0].runs[0].bold = True

                for _, row in content.iterrows():
                    row_cells = docx_table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
                        row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        else:
            st.write(content)
            document.add_paragraph(content)

            
    document = add_images_to_document(image_links,document)
    

    # Save the Word document
    docx_buffer = BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)
try:
    st.download_button("Download Report", data=docx_buffer, file_name=f"valuation_report_{post_id}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    print(image_links)
except: pass
